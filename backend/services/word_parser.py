"""
Word 文档解析服务 - 支持 .docx 文件的解析
"""
import os
from pathlib import Path
from typing import Dict, Optional, List
from docx import Document


def get_mineru_output_dir() -> Path:
    """获取 MinerU 输出目录 - 优先使用环境变量 MINERU_OUTPUT_DIR"""
    output_dir_env = os.getenv('MINERU_OUTPUT_DIR')
    if output_dir_env:
        return Path(output_dir_env)
    
    # 默认使用用户可写目录
    if os.name == 'nt':  # Windows
        base_dir = Path(os.environ.get('LOCALAPPDATA', Path.home() / 'AppData' / 'Local'))
    elif os.uname().sysname == 'Darwin':  # macOS
        base_dir = Path.home() / 'Library' / 'Application Support'
    else:  # Linux
        base_dir = Path(os.environ.get('XDG_DATA_HOME', Path.home() / '.local' / 'share'))
    return base_dir / 'Uverse' / 'outputs'


class WordParser:
    """Word 文档解析器 - 支持 .docx 格式"""
    
    def __init__(self, output_dir: str = None):
        if output_dir is None:
            output_dir = str(get_mineru_output_dir())
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def parse_docx(
        self, 
        file_path: str, 
        doc_id: str,
        progress_callback: Optional[callable] = None,
        log_callback: Optional[callable] = None
    ) -> Dict:
        """
        解析 Word (.docx) 文件
        
        Args:
            file_path: Word 文件路径
            doc_id: 文档 ID
            progress_callback: 进度回调函数
            log_callback: 日志回调函数
            
        Returns:
            解析结果字典
        """
        try:
            if log_callback:
                log_callback("INFO", f"开始解析 Word 文档: {file_path}")
            
            if progress_callback:
                progress_callback(0, 100, "正在读取 Word 文档...")
            
            # 打开 Word 文档
            doc = Document(file_path)
            
            if progress_callback:
                progress_callback(30, 100, "正在提取文本内容...")
            
            # 提取段落文本
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                
                # 每 100 个段落更新一次进度
                if i % 100 == 0 and progress_callback:
                    progress = 30 + int((i / max(len(doc.paragraphs), 1)) * 30)
                    progress_callback(progress, 100, f"已处理 {i} 个段落...")
            
            if log_callback:
                log_callback("INFO", f"提取了 {len(paragraphs)} 个段落")
            
            if progress_callback:
                progress_callback(60, 100, "正在提取表格内容...")
            
            # 提取表格内容
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
                
                if i % 10 == 0 and progress_callback:
                    progress = 60 + int((i / max(len(doc.tables), 1)) * 20)
                    progress_callback(progress, 100, f"已处理 {i} 个表格...")
            
            if log_callback:
                log_callback("INFO", f"提取了 {len(tables)} 个表格")
            
            if progress_callback:
                progress_callback(80, 100, "正在转换为 Markdown...")
            
            # 创建输出目录
            doc_output_dir = self.output_dir / doc_id
            doc_output_dir.mkdir(parents=True, exist_ok=True)
            
            # 转换为 Markdown 格式
            md_content = self._convert_to_markdown(paragraphs, tables)
            
            if progress_callback:
                progress_callback(90, 100, "正在保存 Markdown 文件...")
            
            # 保存 Markdown 文件
            original_name = Path(file_path).stem
            md_filename = f"{original_name}.md"
            md_path = doc_output_dir / md_filename
            
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            if log_callback:
                log_callback("INFO", f"Markdown 文件已保存: {md_path}")
            
            if progress_callback:
                progress_callback(100, 100, "解析完成")
            
            return {
                "doc_id": doc_id,
                "file_type": "docx",
                "total_pages": len(doc.sections) if doc.sections else 1,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
                "output_dir": str(doc_output_dir),
                "markdown_file": str(md_path),
                "images_dir": None,
                "content": md_content,
            }
            
        except Exception as e:
            error_msg = f"解析 Word 文档失败: {str(e)}"
            if log_callback:
                log_callback("ERROR", error_msg)
            raise Exception(error_msg)
    
    def _convert_to_markdown(
        self, 
        paragraphs: List[str], 
        tables: List[List[List[str]]]
    ) -> str:
        """
        将 Word 内容转换为 Markdown 格式
        
        Args:
            paragraphs: 段落列表
            tables: 表格数据列表
            
        Returns:
            Markdown 格式的字符串
        """
        md_lines = []
        
        # 处理段落
        for i, para in enumerate(paragraphs):
            # 检测是否是标题（基于样式和内容的启发式规则）
            md_line = self._parse_paragraph_style(para, i, paragraphs)
            md_lines.append(md_line)
        
        # 处理表格
        for table in tables:
            if table:
                md_lines.append('')
                md_lines.append(self._table_to_markdown(table))
                md_lines.append('')
        
        return '\n\n'.join(md_lines)
    
    def _parse_paragraph_style(self, para: str, index: int, all_paras: List[str]) -> str:
        """
        解析段落样式并转换为 Markdown
        
        使用启发式规则检测标题、列表等格式
        """
        para = para.strip()
        if not para:
            return ''
        
        # 检测一级标题：全大写且较短，或者是文档的第一个段落
        if para.isupper() and len(para) < 100 and len(para) > 5:
            return f"# {para}"
        
        # 检测二级标题：以冒号或特定字符结尾
        if para.endswith(('：', ':')) and len(para) < 80:
            return f"## {para.rstrip('：:').strip()}"
        
        # 检测三级标题：包含序号如 1. 2. 3.
        if para[:2].isdigit() and para[2:3] in '.、':
            return f"### {para}"
        
        # 检测中文序号标题
        if para[:3] in ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、']:
            return f"## {para}"
        
        if para[:4] in ['（一）', '（二）', '（三）', '（四）', '（五）', 
                        '（六）', '（七）', '（八）', '（九）', '（十）']:
            return f"### {para}"
        
        # 检测列表项
        if para.startswith(('•', '-', '*', '·', '○', '●')):
            return f"- {para[1:].strip()}"
        
        # 检测数字列表
        if len(para) > 2 and para[0].isdigit() and para[1:2] in '.、':
            return f"1. {para[2:].strip()}"
        
        # 检测括号数字列表
        if para[:2] in ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩']:
            return f"- {para}"
        
        # 普通段落
        return para
    
    def _table_to_markdown(self, table: List[List[str]]) -> str:
        """
        将表格转换为 Markdown 格式
        
        Args:
            table: 表格数据，每行是一个字符串列表
            
        Returns:
            Markdown 表格格式的字符串
        """
        if not table:
            return ""
        
        md_lines = []
        
        # 获取最大列数
        max_cols = max(len(row) for row in table)
        
        # 表头行
        header = table[0] if table else [''] * max_cols
        # 补齐列数
        header = header + [''] * (max_cols - len(header))
        header = header[:max_cols]
        
        md_lines.append('| ' + ' | '.join(header) + ' |')
        md_lines.append('|' + '|'.join(['---' for _ in range(max_cols)]) + '|')
        
        # 数据行
        for row in table[1:]:
            # 补齐列数
            row = row + [''] * (max_cols - len(row))
            row = row[:max_cols]
            md_lines.append('| ' + ' | '.join(row) + ' |')
        
        return '\n'.join(md_lines)
    
    def get_docx_info(self, file_path: str) -> Dict:
        """
        获取 Word 文档的基本信息（不解析内容）
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            文档信息字典
        """
        try:
            doc = Document(file_path)
            
            # 统计段落数（非空段落）
            paragraph_count = sum(1 for p in doc.paragraphs if p.text.strip())
            
            # 统计表格数
            table_count = len(doc.tables)
            
            # 统计节数（页数估算）
            section_count = len(doc.sections)
            
            return {
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "section_count": section_count,
                "file_size": Path(file_path).stat().st_size,
            }
        except Exception as e:
            return {
                "error": str(e),
                "paragraph_count": 0,
                "table_count": 0,
                "section_count": 0,
            }


# 全局解析器实例
_word_parser = None


def get_word_parser() -> WordParser:
    """获取 Word 解析器实例"""
    global _word_parser
    if _word_parser is None:
        _word_parser = WordParser()
    return _word_parser
